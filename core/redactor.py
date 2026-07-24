"""Document-level redaction pipeline.

Walks a python-docx Document paragraph-by-paragraph AND table-cell-by-cell,
runs all detectors, resolves overlapping spans, applies the PII mapper,
and writes a new .docx preserving all formatting.
"""

from __future__ import annotations

import json
import copy
from pathlib import Path
from typing import List, Tuple, Dict, Any, Optional

from docx import Document
from docx.text.paragraph import Paragraph

from detectors.base import DetectionResult
from detectors.emails import EmailDetector
from detectors.phones import PhoneDetector
from detectors.ssn import SSNDetector
from detectors.credit_card import CreditCardDetector
from detectors.dob import DOBDetector
from detectors.ip_address import IPAddressDetector
from detectors.names import NameDetector
from detectors.companies import CompanyDetector
from detectors.addresses import AddressDetector
from core.mapper import PIIMapper


# ── Overlap resolution ────────────────────────────────────────────────────

def resolve_overlaps(detections: List[DetectionResult]) -> List[DetectionResult]:
    """Remove overlapping spans, keeping the longer / higher-confidence one.

    When two spans overlap, the one with wider coverage wins.  If equal
    width, the one with higher confidence wins.
    """
    if not detections:
        return []

    # Sort by start, then by descending length, then descending confidence
    sorted_dets = sorted(
        detections,
        key=lambda d: (d.start, -(d.end - d.start), -d.confidence),
    )

    resolved: List[DetectionResult] = []
    for det in sorted_dets:
        if resolved and det.start < resolved[-1].end:
            # Overlap — keep the wider span
            prev = resolved[-1]
            prev_len = prev.end - prev.start
            det_len = det.end - det.start
            if det_len > prev_len or (det_len == prev_len and det.confidence > prev.confidence):
                resolved[-1] = det
            # Otherwise keep the previous (already wider/higher confidence)
        else:
            resolved.append(det)

    return resolved


# ── Run-level text replacement (format-preserving) ────────────────────────

def _replace_in_paragraph(paragraph: Paragraph, replacements: List[Tuple[int, int, str]]) -> None:
    """Apply replacements to a paragraph while preserving Run formatting.

    *replacements* is a list of ``(start, end, new_text)`` sorted by start
    descending (so we replace from right to left to keep offsets valid).
    """
    if not replacements or not paragraph.runs:
        return

    # Build a map of (char_offset → run_index, offset_within_run)
    runs = paragraph.runs
    run_texts = [r.text for r in runs]

    # Apply replacements right-to-left so earlier offsets stay valid
    for start, end, new_text in sorted(replacements, key=lambda r: -r[0]):
        _apply_single_replacement(runs, run_texts, start, end, new_text)

    # Write back
    for i, run in enumerate(runs):
        run.text = run_texts[i]


def _apply_single_replacement(
    runs, run_texts: List[str], start: int, end: int, new_text: str
) -> None:
    """Replace characters [start, end) across possibly multiple runs."""
    # Find which runs contain the start and end offsets
    offset = 0
    start_run = start_off = end_run = end_off = -1

    for i, txt in enumerate(run_texts):
        run_start = offset
        run_end = offset + len(txt)

        if start_run == -1 and start < run_end:
            start_run = i
            start_off = start - run_start

        if end <= run_end:
            end_run = i
            end_off = end - run_start
            break

        offset += len(txt)
    else:
        # end goes beyond all runs — clamp to last run
        if start_run != -1:
            end_run = len(run_texts) - 1
            end_off = len(run_texts[end_run])

    if start_run == -1 or end_run == -1:
        return  # safety: span outside paragraph text

    if start_run == end_run:
        # Simple case — replacement is within a single run
        txt = run_texts[start_run]
        run_texts[start_run] = txt[:start_off] + new_text + txt[end_off:]
    else:
        # Multi-run replacement: put the new text in the first run,
        # clear intermediate runs, trim the last run
        run_texts[start_run] = run_texts[start_run][:start_off] + new_text
        for mid in range(start_run + 1, end_run):
            run_texts[mid] = ""
        run_texts[end_run] = run_texts[end_run][end_off:]


# ── Main pipeline ─────────────────────────────────────────────────────────

class RedactionPipeline:
    """End-to-end .docx redaction pipeline."""

    def __init__(self, mapping_path: Optional[str] = None):
        self.mapper = PIIMapper(mapping_path=mapping_path)
        self.detectors = [
            EmailDetector(),
            PhoneDetector(),
            SSNDetector(),
            CreditCardDetector(),
            DOBDetector(),
            IPAddressDetector(),
            NameDetector(),
            CompanyDetector(),
            AddressDetector(),
        ]
        self.redaction_log: List[Dict[str, Any]] = []
        self._stats: Dict[str, int] = {}

    def process(
        self,
        input_path: str,
        output_path: str,
        mapping_output: Optional[str] = None,
        log_output: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Run the full pipeline on a .docx file.

        Returns a summary dict with counts by PII type.
        """
        doc = Document(input_path)
        self.redaction_log = []
        self._stats = {}

        # Process all paragraphs (body)
        for para in doc.paragraphs:
            self._process_paragraph(para, "body")

        # Process all table cells
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        self._process_paragraph(
                            para,
                            f"table[{table_idx}].row[{row_idx}].cell[{cell_idx}]",
                        )

        # Save outputs
        doc.save(output_path)

        if mapping_output:
            self.mapper.save(mapping_output)

        if log_output:
            with open(log_output, "w", encoding="utf-8") as fh:
                json.dump(self.redaction_log, fh, indent=2, ensure_ascii=False)

        return self._build_summary()

    def process_bytes(self, file_bytes, mapping_output=None, log_output=None):
        """Process a .docx from bytes (for API use). Returns (output_bytes, summary)."""
        import io
        input_stream = io.BytesIO(file_bytes)
        doc = Document(input_stream)
        self.redaction_log = []
        self._stats = {}

        for para in doc.paragraphs:
            self._process_paragraph(para, "body")

        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        self._process_paragraph(
                            para,
                            f"table[{table_idx}].row[{row_idx}].cell[{cell_idx}]",
                        )

        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)

        if mapping_output:
            self.mapper.save(mapping_output)
        if log_output:
            with open(log_output, "w", encoding="utf-8") as fh:
                json.dump(self.redaction_log, fh, indent=2, ensure_ascii=False)

        return output_stream.read(), self._build_summary()

    # ── internal helpers ──────────────────────────────────────────────────

    def _process_paragraph(self, para: Paragraph, location: str) -> None:
        """Detect and replace PII in a single paragraph."""
        full_text = para.text
        if not full_text or not full_text.strip():
            return

        # Run all detectors
        all_detections: List[DetectionResult] = []
        for detector in self.detectors:
            try:
                all_detections.extend(detector.detect(full_text))
            except Exception:
                continue  # detector failures should not crash the pipeline

        if not all_detections:
            return

        # Resolve overlaps
        resolved = resolve_overlaps(all_detections)

        # Build replacements
        replacements: List[Tuple[int, int, str]] = []
        for det in resolved:
            fake_value = self.mapper.get_or_create(det.text, det.pii_type)
            replacements.append((det.start, det.end, fake_value))

            # Update stats
            self._stats[det.pii_type] = self._stats.get(det.pii_type, 0) + 1

            # Audit log (no original values in log — production safe)
            self.redaction_log.append({
                "location": location,
                "pii_type": det.pii_type,
                "confidence": det.confidence,
                "start": det.start,
                "end": det.end,
                "replacement": fake_value,
            })

        # Apply to the paragraph's runs
        _replace_in_paragraph(para, replacements)

    def _build_summary(self) -> Dict[str, Any]:
        """Compile the final summary report."""
        return {
            "total_redactions": sum(self._stats.values()),
            "by_type": dict(self._stats),
            "unique_mappings": self.mapper.get_stats(),
            "total_mapping_uses": self.mapper.get_total_replacements(),
        }
