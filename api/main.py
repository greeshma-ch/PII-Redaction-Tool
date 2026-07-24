"""FastAPI application for the PII Redaction service.

Endpoints:
    GET  /           — Serve the frontend UI
    GET  /health      — Health check for Render
    POST /redact      — Upload a .docx, receive redacted .docx + JSON diff
    GET  /download/{filename} — Download a redacted file
"""

import os
import time
import tempfile
from pathlib import Path

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document

from core.redactor import RedactionPipeline

app = FastAPI(title="PII Redaction Tool")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

TEMP_DIR = Path(tempfile.mkdtemp(prefix="pii_redactor_"))
FRONTEND_DIR = Path(__file__).resolve().parent.parent / "frontend"


def _extract_text(doc_path: str, limit: int = 8000) -> str:
    """Pull plain text from a .docx for the diff preview."""
    doc = Document(doc_path)
    parts = []
    total = 0
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
            total += len(t)
            if total >= limit:
                break
    if total < limit:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        parts.append(t)
                        total += len(t)
                        if total >= limit:
                            break
                if total >= limit:
                    break
            if total >= limit:
                break
    return "\n".join(parts)[:limit]


# ── Routes ────────────────────────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def root():
    index = FRONTEND_DIR / "index.html"
    if index.exists():
        return HTMLResponse(index.read_text(encoding="utf-8"))
    return HTMLResponse("<h1>Frontend not found</h1>", status_code=404)


@app.get("/health")
async def health():
    return {"status": "ok"}


@app.post("/redact")
async def redact(file: UploadFile = File(...)):
    """Accept a .docx upload, redact PII, return JSON summary + download link."""
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "Only .docx files are supported.")

    content = await file.read()
    if len(content) > 50 * 1024 * 1024:
        raise HTTPException(400, "File exceeds the 50 MB limit.")

    ts = int(time.time() * 1000)
    input_path = TEMP_DIR / f"upload_{ts}_{file.filename}"
    output_name = f"redacted_{ts}_{file.filename}"
    output_path = TEMP_DIR / output_name

    input_path.write_bytes(content)

    try:
        t0 = time.time()
        pipeline = RedactionPipeline()
        summary = pipeline.process(
            input_path=str(input_path),
            output_path=str(output_path),
        )
        elapsed = time.time() - t0
        print(f"[redact] {file.filename}: {summary['total_redactions']} redactions in {elapsed:.1f}s")

        original_text = _extract_text(str(input_path))
        redacted_text = _extract_text(str(output_path))

        return JSONResponse({
            "summary": summary,
            "redacted_file_url": f"/download/{output_name}",
            "original_text_preview": original_text,
            "redacted_text_preview": redacted_text,
            "redaction_log": pipeline.redaction_log[:500],  # Cap for response size
            "processing_time_s": round(elapsed, 2),
        })
    except Exception as exc:
        raise HTTPException(500, f"Redaction failed: {exc}")
    finally:
        if input_path.exists():
            input_path.unlink()


@app.get("/download/{filename}")
async def download(filename: str):
    path = TEMP_DIR / filename
    if not path.exists():
        raise HTTPException(404, "File not found.")
    return FileResponse(
        str(path),
        filename=filename.split("_", 2)[-1] if "_" in filename else filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
